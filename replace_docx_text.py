"""Replace text in all .docx files in the current directory."""
import glob
import re
from docx import Document
from docx.oxml.ns import qn

FIND = "March 29"
REPLACE = "March 30"


def replace_in_paragraph(paragraph):
    """Replace text across runs in a paragraph, preserving formatting."""
    # First try simple per-run replacement
    for run in paragraph.runs:
        if FIND in run.text:
            run.text = run.text.replace(FIND, REPLACE)

    # Handle cases where the text is split across runs
    full_text = "".join(r.text for r in paragraph.runs)
    if FIND in full_text:
        # Rebuild runs: put all text in first run, clear the rest
        new_text = full_text.replace(FIND, REPLACE)
        if paragraph.runs:
            paragraph.runs[0].text = new_text
            for run in paragraph.runs[1:]:
                run.text = ""


def process_doc(path):
    doc = Document(path)
    changed = False

    for para in doc.paragraphs:
        before = "".join(r.text for r in para.runs)
        replace_in_paragraph(para)
        after = "".join(r.text for r in para.runs)
        if before != after:
            changed = True

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    before = "".join(r.text for r in para.runs)
                    replace_in_paragraph(para)
                    after = "".join(r.text for r in para.runs)
                    if before != after:
                        changed = True

    if changed:
        doc.save(path)
        print(f"  Updated: {path}")
    else:
        print(f"  No match: {path}")

    return changed


def main():
    files = glob.glob("/home/user/just-some-stuff/**/*.docx", recursive=True)
    files += glob.glob("/home/user/just-some-stuff/*.docx")
    files = list(set(files))

    if not files:
        print("No .docx files found.")
        return

    print(f"Found {len(files)} .docx file(s). Replacing '{FIND}' -> '{REPLACE}'...\n")
    updated = sum(process_doc(f) for f in files)
    print(f"\nDone. {updated}/{len(files)} file(s) updated.")


if __name__ == "__main__":
    main()
